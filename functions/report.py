from os import path, listdir, makedirs
from typing import List

from cv2 import imread
import math
import json
import numpy as np
from docx import Document
from docx.shared import Inches
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from functions import curve
from functions.slicing import getPath

DIRDCT=getPath()
PSIZE=2.5

def calLine(m,n):
    d=(m[0]-n[0])**2+(m[1]-n[1])**2
    return d

def calDiameter(maskfile,npyfile,w):
    img = imread(maskfile, 0)
    ds=list(set(zip(np.nonzero(img)[0],np.nonzero(img)[1])))
    
    maxline=0
    for m in ds:
        for n in ds:
            d=calLine(m,n)
            maxline=max(d,maxline)
            
    npy_data = np.load(npyfile)
    
    t,hu_all=0,0
    
    for x,y in ds:
        t += 1
        x1 = int((x-118) * 1024 / 788)
        y1 = int((y-118) * 1024 / 788)

        hu_all += npy_data[x1][y1]
    return round(math.sqrt(maxline)*w,2),round(hu_all/t,2)

def getMeta(mhdPath):
    with open(mhdPath) as f:
        t = f.read().split("\n")[:-1]
    di=[]
    dz=[]
    for i in range(len(t)):
        st=t[i].split(" = ")
        di.append(st[0])
        dz.append(st[1])
    mhddata=dict(zip(di,dz))
    meta={}
    meta["name"]=mhddata["0010|0010"]
    meta["age"]=mhddata["0010|1010"]
    meta["gender"]=mhddata["0010|0040"]
    meta["date"]=mhddata["0008|0012"]
    meta["spacing"]=float(mhddata["ElementSpacing"].split(" ")[0])
    return meta

def addMeta(document,meta):
    document.add_paragraph('姓名:%s'%meta["name"])
    document.add_paragraph('年龄:%s'%meta["age"])
    document.add_paragraph('性别:%s'%meta["gender"])

def addCompareImage(document,patient,dates,caseImages):
    p = document.add_paragraph()
    #tit.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    r = p.add_run()
    for t in dates:
        r.add_text("%-15s%-15s"%(str(t)+" raw",str(t)+" rst"))\

    l = 0
    for k in list(caseImages.keys()):
        if l > len(caseImages[k]):
            pass
        else:
            l = len(caseImages[k])

    for n in range(l):
        p = document.add_paragraph()
        #p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r = p.add_run()
        for c in list(caseImages.keys()):
            if n<len(caseImages[c]):
                raw = DIRDCT["mhdpng"]%(patient,c)
                bound = DIRDCT["bound"]%(patient,c)
                r.add_picture(raw+"/"+caseImages[c][n]+".png", width=Inches(PSIZE))
                r.add_picture(bound+"/"+caseImages[c][n]+".png", width=Inches(PSIZE))
            else:
                r.add_text(" "*30)

def addVolume(document,vs,ms,ds,ts):
    p = document.add_paragraph()
    r=p.add_run()
    r.add_text("%-17s%-17s%-17s%-17s"%("date","volume(mm**3))","mass(mg)","density(mg/mm*3)"))
    for i in range(len(ts)):
        p = document.add_paragraph()
        r=p.add_run()
        r.add_text("%-17s%-17s%-17s%-17s"%(str(ts[i]),str(round(vs[i],2)),str(round(ms[i],2)),str(round(ds[i],2))))

    p = document.add_paragraph()
    p = document.add_paragraph()
    p = document.add_paragraph()

    if not path.exists("temp"):
        makedirs("temp")

    if len(ts)>1:
        curve.vdtCurve(ts,vs,"temp/vdt.png")
        curve.mdtCurve(ts,ms,"temp/wdt.png")
        curve.ddtCurve(ts,ds,"temp/ddt.png")

        p = document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        r=p.add_run()
        r.add_picture("temp/vdt.png", width=Inches(PSIZE*2))
        r.add_picture("temp/wdt.png", width=Inches(PSIZE*2))
        r.add_picture("temp/ddt.png", width=Inches(PSIZE*2))

def addDiameter(document,patient,caseImages,meta):
    for c in list(caseImages.keys()):
        p = document.add_paragraph()
        r = p.add_run()
        dias,dsts = [],[]
        for name in caseImages[c]:
            maskfile = DIRDCT["rstMask"]%(patient,c)+"/"+name+".png"
            npyfile = DIRDCT["mhdnpy"]%(patient,c)+"/"+name+".npy"
            dia,dst = calDiameter(maskfile,npyfile,meta["spacing"])
            dias.append(dia)
            dsts.append(dst)
        raw = DIRDCT["mhdpng"]%(patient,c)
        bound = DIRDCT["bound"]%(patient,c)
        r.add_text(c+" 最大直径的截面")
        loc = dias.index(max(dias))
        p = document.add_paragraph()
        r = p.add_run()
        r.add_picture(raw+"/"+caseImages[c][loc]+".png", width=Inches(PSIZE))
        r.add_picture(bound+"/"+caseImages[c][loc]+".png", width=Inches(PSIZE))
        p = document.add_paragraph()
        r = p.add_run()
        r.add_text("%s 最大直径为: %.2f mm,截面密度: %.2f HU"%(caseImages[c][loc]+".png",dias[loc],dsts[loc]))
        
        p = document.add_paragraph()
        r = p.add_run()
        r.add_text(c+" 密度最大的截面")
        loc = dsts.index(max(dsts))
        p = document.add_paragraph()
        r = p.add_run()
        r.add_picture(raw+"/"+caseImages[c][loc]+".png", width=Inches(PSIZE))
        r.add_picture(bound+"/"+caseImages[c][loc]+".png", width=Inches(PSIZE))
        p = document.add_paragraph()
        r = p.add_run()
        r.add_text("%s 最大直径为: %.2f mm,截面密度: %.2f HU"%(caseImages[c][loc]+".png",dias[loc],dsts[loc]))
        
        p = document.add_paragraph()
        p = document.add_paragraph()


def addDiameter2(document, l_diameter,l_diameter_data,l_density,l_density_data):
    p = document.add_paragraph()
    r = p.add_run()
    for i in range(len(l_diameter)):
        st = l_diameter[i].split("/")
        patient,case,pname = st[1], st[2], st[-1]
        bound_path = l_diameter[i].replace("mhdpng","bound")

        r.add_text(patient + " " + case + " 最大直径的截面")
        p = document.add_paragraph()
        r = p.add_run()
        r.add_picture(l_diameter[i], width=Inches(PSIZE))
        r.add_picture(bound_path, width=Inches(PSIZE))
        p = document.add_paragraph()
        r = p.add_run()
        r.add_text(pname + " 最大直径为: %.2f mm,截面密度: %.2f HU" % (l_diameter_data[i][0], l_diameter_data[i][1]))

        st = l_density[i].split("/")
        patient, case, pname = st[1], st[2], st[-1]
        bound_path = l_density[i].replace("mhdpng", "bound")
        p = document.add_paragraph()
        r = p.add_run()
        r.add_text(patient + " " + case + " 密度最大的截面")
        p = document.add_paragraph()
        r = p.add_run()
        r.add_picture(l_density[i], width=Inches(PSIZE))
        r.add_picture(bound_path, width=Inches(PSIZE))
        p = document.add_paragraph()
        r = p.add_run()
        r.add_text(pname + " 最大直径为: %.2f mm,截面密度: %.2f HU" % (l_density_data[i][0], l_density_data[i][1]))


def addReport(patient):
    document = Document()
    document.styles['Normal'].font.name = u'宋体'
    document.styles['Normal'].font.size=Pt(24)
    document.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
    sec=document.sections[-1]
    sec.page_height,sec.page_width=10058400*3, 7772400*3

    casesDir = "project/%s"%patient
    cases = [c for c in listdir(casesDir) if path.isdir(casesDir+"/"+c)]

    mhds = [DIRDCT["rawData"]%(patient,c)+"/%s.mhd"%c for c in cases]
    meta = getMeta(mhds[-1])
    addMeta(document,meta)

    document.add_paragraph()
    document.add_paragraph()

    vs,ms,ds,ts=[],[],[],[]
    l_diameter_data: List[List[float]]
    l_density_data: List[List[float]]
    l_diameter,l_diameter_data, l_density, l_density_data= [],[],[],[]


    try:
        jsons = []
        for c in cases:
            for f in listdir(DIRDCT["caseDir"]%(patient,c)):
                if f.endswith(".json"):
                    jsons.append(DIRDCT["caseDir"]%(patient,c)+"/"+f)
    except:
        print("可能未生成结果文件")
        print("-"*20)

    try:
        for j in jsons:
            with open(j,"r") as f:
                r = json.load(f)
            vs.append(float(r["volume"].replace(" mm**3","")))
            ms.append(float(r["mass"].replace(" mg","")))
            ds.append(float(r["density"].replace(" mg/mm**3","")))

    except:
        print("可能未计算体积")
        print("-"*20)

    try:
        for j in jsons:
            with open(j, "r") as f:
                r = json.load(f)
            l_diameter.append(r["longest_diameter"])
            l_diameter_data.append([float(r["longest_diameter_diameter"].split(" ")[0]), float(r["longest_diameter_density"].split(" ")[0])])
            l_density.append(r["largest_density"])
            l_density_data.append([float(r["largest_density_diameter"].split(" ")[0]), float(r["largest_density_density"].split(" ")[0])])
    except:
        print("可能未计算最大直径和最大密度")
        print("-" * 20)


    for mhd in mhds:
        ts.append(getMeta(mhd)["date"])

            
    caseImage = {}
    for c in cases:
        for j in jsons:
            if not c in j:
                continue
            with open(j,"r") as f:
                r = json.load(f)
            caseImage[c] = [f.split("/")[-1].split(".")[0] for f in r["resultPath"].split(",")]

    addCompareImage(document,patient,ts,caseImage)

    document.add_paragraph()
    document.add_paragraph()

    addVolume(document,vs,ms,ds,ts)

    document.add_paragraph()
    document.add_paragraph()

    addDiameter2(document, l_diameter,l_diameter_data,l_density,l_density_data)
    # addDiameter(document,patient,caseImage,meta)

    reportSave = casesDir+"/%s.docx"%patient
    document.save(reportSave)


if __name__ == "__main__":
    addReport("yeguolian")